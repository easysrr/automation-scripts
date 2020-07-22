import os
import wikipedia
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
par = doc.add_paragraph()
run = doc.add_paragraph().add_run()
blank_par = doc.add_paragraph('')
font = run.font
os.system('cls')

def sysClear():
    os.system('cls')
    welcome()
    print('Enter the subject: ' + key_word)
    print('Enter your name: ' + usr_name)
    print('=============== ===============')
    print('The total number of points will be: ' + str(usr_bulletPointList_input + 1))
    print('')
    

# welcome message
def welcome():
    print(
        """
        -------------------------

                yeet
				
					~Bartosz Rydzewski
        -------------------------
                                    """
    )


# variables
sources_list = []
i = 0
x = 0
src_par = ''
bpl_par = ''
bpl_point = ''
no = 2
ent = ''

# user input
welcome()
key_word = input('Enter the subject: ')
usr_name = input('Enter your name: ')
usr_name = usr_name + ' 2d(sp)'
file_name = key_word + ' - ' + usr_name
wikipedia.set_lang('pl')

try:
    wikipedia_output = wikipedia.summary(key_word)
    sources_list.append(wikipedia.page(key_word).url)
except wikipedia.PageError:
    print('error: Wikipedia matches not found.')
except TimeoutError or wikipedia.HTTPTimeoutError:
    print('error: Connection with Wikipedia could not be established.')
except wikipedia.exceptions.DisambiguationError as e:
    print('error: Multiple wikipedia results. Choosing random')   
    s = random.choice(e.options) 
    wikipedia_output = wikipedia.summary(s)
print('=============== ===============')


# doc template
doc.add_heading(key_word, 0)
usr_sign = doc.add_paragraph()
usr_sign.add_run(usr_name).italic = True
usr_sign.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def_point = doc.add_paragraph('', style = 'List Number')
def_point.add_run('Definicja').font.size = Pt(20)
def_par = doc.add_paragraph(wikipedia_output)


# bullet point list input
usr_bulletPointList_input = int(input('How many points do you want to add? (first point is gonna be the definition from wikipedia): '))
for x in range(0, usr_bulletPointList_input):
    sysClear()
    if x == 0:
        ent = 'nd'
    elif x == 1:
        ent = 'rd'
    else:
        ent = 'th'
    
    bpl_point = doc.add_paragraph('', style = 'List Number')
    bpl_point.add_run(input('What is gonna be the name of the '+ str(no) + ent + ' point?: ')).font.size = Pt(20)
    bpl_par = doc.add_paragraph(input('Please fill out the '+ str(no) + ent + ' point: '))
    sources_list.append(input('Whats the URL of the website?: '))
    sysClear()
    no = no + 1
    x += 1

# sources
doc.add_paragraph(' ')
doc.add_paragraph(' ')
doc.add_paragraph(' ')

doc_sources_list = doc.add_paragraph().add_run('Źródła: ')
doc_sources_list.font.color.rgb = RGBColor(0xA3, 0xA3, 0xA3)
doc_sources_list.italic = True

print(sources_list)
while i != len(sources_list):
    src_par = doc.add_paragraph().add_run(sources_list[i])
    src_par.italic = True
    src_par.font.color.rgb = RGBColor(0xA3, 0xA3, 0xA3)
    i += 1

# program end
doc.save(file_name + '.docx')
print('')
print('Done.')
